from langchain_core.tools import tool
import os
import re
import uuid
import logging
import unicodedata
import asyncio
from typing import Any
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from src.config import settings

logger = logging.getLogger(__name__)

def _safe_filename(title: str, max_len: int = 50) -> str:
    """Normaliza para ASCII puro — evita surrogates e erros de encoding."""
    normalized = unicodedata.normalize('NFKD', title)
    ascii_only = normalized.encode('ascii', errors='ignore').decode('ascii')
    safe = re.sub(r'[^\w\s-]', '', ascii_only).strip().lower().replace(' ', '-')
    return (safe or 'documento')[:max_len]


# ─── Paleta Executiva ────────────────────────────────────────────────────────
_AZUL_CORP  = RGBColor(28,  78, 158)   # Azul corporativo profundo
_AZUL_SEC   = RGBColor(36, 110, 185)   # Azul secundário
_CINZA_H    = RGBColor(45,  45,  45)   # Quase preto para sub-headings
_CINZA_BODY = RGBColor(55,  55,  55)   # Cinza escuro para corpo


# ─── Helpers DOCX ────────────────────────────────────────────────────────────

def _add_rich_text(paragraph, text: str, size_pt: int = 11, color=None):
    """Adiciona texto com suporte a **negrito** inline."""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        bold = part.startswith('**') and part.endswith('**')
        run = paragraph.add_run(part[2:-2] if bold else part)
        run.bold = bold
        run.font.size = Pt(size_pt)
        run.font.name = "Calibri"
        if color:
            run.font.color.rgb = color


def _set_margins(doc: Document, margin_in: float = 1.0):
    for section in doc.sections:
        section.top_margin    = Inches(margin_in)
        section.bottom_margin = Inches(margin_in)
        section.left_margin   = Inches(margin_in)
        section.right_margin  = Inches(margin_in)


def _parse_markdown_to_docx(doc: Document, content: str):
    """Converte markdown simples em Word com design executivo."""
    for line in content.split('\n'):
        stripped = line.strip()

        if not stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            continue

        if stripped.startswith('# ') and not stripped.startswith('## '):
            h = doc.add_heading(stripped[2:].strip(), level=1)
            h.paragraph_format.space_before = Pt(20)
            h.paragraph_format.space_after  = Pt(8)
            if h.runs:
                h.runs[0].font.color.rgb = _AZUL_CORP
                h.runs[0].font.size      = Pt(16)
                h.runs[0].font.name      = "Calibri"

        elif stripped.startswith('## '):
            h = doc.add_heading(stripped[3:].strip(), level=2)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after  = Pt(6)
            if h.runs:
                h.runs[0].font.color.rgb = _AZUL_SEC
                h.runs[0].font.size      = Pt(13)
                h.runs[0].font.name      = "Calibri"

        elif stripped.startswith('### '):
            h = doc.add_heading(stripped[4:].strip(), level=3)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after  = Pt(4)
            if h.runs:
                h.runs[0].font.color.rgb = _CINZA_H
                h.runs[0].font.size      = Pt(11)
                h.runs[0].font.name      = "Calibri"

        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.space_before = Pt(2)
            _add_rich_text(p, stripped[2:].strip(), size_pt=11, color=_CINZA_BODY)

        elif re.match(r'^\d+\. ', stripped):
            text_body = re.sub(r'^\d+\. ', '', stripped)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(4)
            _add_rich_text(p, text_body, size_pt=11, color=_CINZA_BODY)

        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(8)
            p.paragraph_format.space_before = Pt(2)
            _add_rich_text(p, stripped, size_pt=11, color=_CINZA_BODY)


from pydantic import BaseModel, Field
from typing import Optional, Any, Union

class DocxSchema(BaseModel):
    title: Optional[str] = Field(default="Documento", description="Título do documento.")
    content: Optional[str] = Field(default=None, description="Conteúdo em texto ou markdown para o documento.")
    filename: Optional[str] = Field(default=None, description="Nome do arquivo .docx opcional.")

class PdfSchema(BaseModel):
    title: Optional[str] = Field(default="Documento", description="Título do PDF.")
    content: Optional[str] = Field(default=None, description="Conteúdo do PDF.")

@tool(args_schema=DocxSchema)
async def generate_docx(title: str = "Documento", content: str = "", filename: Optional[str] = None) -> str:
    """Gera DOCX com tratamento robusto e formatação melhorada."""
    if not content:
        logger.error("[DOCXGen] ERRO CRÍTICO: content é None ou vazio!")
        return "Erro: O conteúdo do documento não pode estar vazio."

    try:
        # ==================================================================
        # NORMALIZAÇÃO DE PARÂMETROS
        # ==================================================================
        if not filename:
            filename = f"documento_{int(datetime.now().timestamp())}.docx"
        
        if not filename.endswith(".docx"):
            filename += ".docx"
        output_dir = settings.DATA_OUTPUTS_PATH
        os.makedirs(output_dir, exist_ok=True)
        filepath = os.path.join(output_dir, filename)
        
        logger.info(f"[DOCXGen] Gerando DOCX: {filepath}")
        
        # ==================================================================
        # FUNÇÃO SÍNCRONA PARA RODAR EM THREAD SEPARADA
        # ==================================================================
        def _generate():
            doc = Document()
            
            # Configura margens
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # TÍTULO PRINCIPAL
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # DATA
            date_para = doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # LINHA HORIZONTAL
            doc.add_paragraph('_' * 50)
            
            # CONTEÚDO COM ESTILO CORPORATIVO PREMIUM
            _parse_markdown_to_docx(doc, str(content))
            
            # RODAPÉ
            doc.add_paragraph('_' * 50)
            footer = doc.add_paragraph("Documento gerado pelo Arth Executive")
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer.paragraph_format.space_before = Pt(12)
            
            doc.save(filepath)
            return filepath
        
        # Executa em thread separada
        filepath = await asyncio.to_thread(_generate)
        
        # Verifica resultado
        if os.path.exists(filepath) and os.path.getsize(filepath) > 0:
            size_bytes = os.path.getsize(filepath)
            logger.info(f"[DOCXGen] ✅ DOCX gerado: {filepath} ({size_bytes} bytes)")
            return f"Documento Word gerado com sucesso: <SEND_FILE:{os.path.basename(filepath)}>"
        else:
            logger.error(f"[DOCXGen] ❌ Arquivo não foi criado")
            return "Falha ao gerar DOCX: Arquivo não foi criado."
            
    except Exception as e:
        logger.error(f"[DOCXGen] ❌ Erro: {str(e)}")
        return f"Erro ao gerar DOCX: {str(e)}"


# ─── PDF ─────────────────────────────────────────────────────────────────────

# ─── Paleta Executiva Manus AI ───────────────────────────────────────────────
_NAVY   = (10, 12, 16)      # Navy Profundo
_COBALT = (88, 166, 255)    # Azul Cobalto (Electric)
_TEXT   = (50, 50, 50)      # Cinza Escuro
_AZUL_CORP_PDF = (28, 78, 158)  # Azul Corporativo (Tupla para PDF)

class ArthPDF(FPDF):
    def header(self):
        # Logo/Marca d'água superior
        self.set_font("Helvetica", "B", 10)
        self.set_text_color(120, 120, 120)
        self.cell(0, 10, "ARTH EXECUTIVE  ·  INTELIGÊNCIA ESTRATÉGICA", 0, 1, "R")
        
        # Barra decorativa Cobalto
        self.set_draw_color(*_COBALT)
        self.set_line_width(0.5)
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(10)

    def footer(self):
        self.set_y(-15)
        # Linha superior do rodapé
        self.set_draw_color(200, 200, 200)
        self.line(10, self.get_y(), 200, self.get_y())
        
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(150, 150, 150)
        date_str = datetime.now().strftime("%d/%m/%Y %H:%M")
        self.cell(0, 10, f"Gerado por Arth Executive em {date_str}  |  Confidencial  |  Página {self.page_no()}", 0, 0, "C")

def _clean_pdf_text(text: str) -> str:
    """Limpa texto para compatibilidade total com FPDF Latin-1."""
    if not text: return ""
    # Substitui caracteres comuns de interrupção
    text = text.replace('–', '-').replace('—', '-').replace('‘', "'").replace('’', "'").replace('“', '"').replace('”', '"')
    # Força codificação latin-1 ignorando o que não couber
    return text.encode("latin-1", errors="replace").decode("latin-1")

def _safe_multi_cell(pdf, w, h, txt):
    """Fallback seguro para textos sem espaços (como URLs longas) que quebram o FPDF."""
    try:
        pdf.multi_cell(w, h, txt)
    except Exception as e:
        # Palavras gigantes sem espaço causam "Not enough horizontal space"
        # Quebramos a força
        import textwrap
        wrapped = "\\n".join(textwrap.wrap(txt, width=90, break_long_words=True))
        try:
            pdf.multi_cell(w, h, wrapped)
        except:
            pass # se ainda der erro, ignora a linha fatal

@tool(args_schema=PdfSchema)
async def generate_pdf(title: str, content: str) -> str:
    """Cria um documento PDF Executivo Premium com design Manus AI."""
    if title is None or content is None:
        logger.error(f"[PDFGen] ERRO: Parâmetros nulos. title={title}")
        title = title or "Documento Executivo"
        content = content or "Conteúdo não fornecido."
        
    try:
        # Garante nome de arquivo limpo e único
        clean_title = _safe_filename(title)
        filename = f"{uuid.uuid4().hex[:6]}-{clean_title}.pdf"
        output_dir = os.path.abspath(settings.DATA_OUTPUTS_PATH)
        os.makedirs(output_dir, exist_ok=True)
        filepath = os.path.join(output_dir, filename)

        pdf = ArthPDF()
        pdf.set_margins(left=15, top=15, right=15)
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.add_page()

        # TÍTULO PRINCIPAL (Box Colorido)
        pdf.set_fill_color(*_AZUL_CORP_PDF)
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("Helvetica", "B", 20)
        pdf.cell(0, 15, _clean_pdf_text(title.upper()), 0, 1, 'C', fill=True)
        pdf.ln(10)

        # Configurações de largura efetiva
        eff_w = pdf.w - pdf.l_margin - pdf.r_margin
        pdf.set_text_color(*_TEXT)

        # Processamento de conteúdo com suporte a Markdown básico
        for line in content.split('\n'):
            stripped = line.strip()
            if not stripped:
                pdf.ln(4)
                continue

            clean_line = _clean_pdf_text(stripped)

            if stripped.startswith('# ') and not stripped.startswith('## '):
                pdf.set_font("Helvetica", "B", 16)
                pdf.set_text_color(*_AZUL_CORP_PDF)
                _safe_multi_cell(pdf, eff_w, 10, clean_line[2:].replace("**", ""))
                pdf.ln(2)
            elif stripped.startswith('## '):
                pdf.set_font("Helvetica", "B", 14)
                pdf.set_text_color(*_AZUL_CORP_PDF)
                _safe_multi_cell(pdf, eff_w, 9, clean_line[3:].replace("**", ""))
                pdf.ln(1)
            elif stripped.startswith('- ') or stripped.startswith('* '):
                pdf.set_font("Helvetica", "", 11)
                pdf.set_text_color(*_TEXT)
                _safe_multi_cell(pdf, eff_w, 7, f"  • {clean_line[2:]}")
            else:
                pdf.set_font("Helvetica", "", 11)
                pdf.set_text_color(*_TEXT)
                # Formata texto em negrito básico para PDF simples
                if "**" in clean_line:
                    clean_line = clean_line.replace("**", "")
                    pdf.set_font("Helvetica", "B", 11)
                _safe_multi_cell(pdf, eff_w, 7, clean_line)
                pdf.ln(3) # Espaçamento elegante entre parágrafos

        await asyncio.to_thread(pdf.output, filepath)
        
        if os.path.exists(filepath):
            logger.info(f"[PDF] ✅ PDF Premium salvo: {filepath} ({os.path.getsize(filepath)} bytes)")
            return f"PDF Executivo gerado com sucesso: <SEND_FILE:{filename}>"
        return "Falha ao gravar arquivo PDF."
    except Exception as e:
        logger.error(f"[PDF] ❌ Erro: {e}", exc_info=True)
        return f"Falha no PDF: {str(e)}"
    except Exception as e:
        logger.error(f"[PDF] Erro: {e}", exc_info=True)
        return f"Falha ao gerar PDF: {str(e)}"
